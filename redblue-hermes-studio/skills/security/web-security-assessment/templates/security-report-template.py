#!/usr/bin/env python3
"""
Security Assessment Report Generator (DOCX)
Usage: python3 /path/to/security-report-template.py
Requires: pip install python-docx
Edit the configuration section at top, then run.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

# === Configuration ===
TARGET_DOMAIN = "example.com"
TARGET_NAME = "Example Corp Website"
ASSESSMENT_DATE = datetime.datetime.now().strftime('%Y年%m月%d日')
OUTPUT_PATH = f'/tmp/{TARGET_DOMAIN}_安全评估报告.docx'

FINDINGS = [
    # (severity, title, description, evidence, impact, remediation, priority)
    # severity: 'critical'|'high'|'medium'|'low'
    # priority: 'P0 紧急'|'P1 高'|'P2 中'|'P3 低'
]

TARGET_INFO = [
    # ('项目', '详情')
]
# =====================

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

RISK_COLORS = {'critical': 'CC0000', 'high': 'E67E00', 'medium': 'D4A017', 'low': '2E8B57'}
RISK_LABELS = {'critical': '严重', 'high': '高危', 'medium': '中危', 'low': '低危'}

def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def make_table(doc, headers, rows, header_color="2B579A"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = '微软雅黑'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        set_cell_shading(cell, header_color)
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9.5)
                    run.font.name = '微软雅黑'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            if ri % 2 == 1:
                set_cell_shading(cell, "F2F2F2")
    return table

def add_risk_badge(paragraph, severity):
    color = RISK_COLORS.get(severity, '666666')
    label = RISK_LABELS.get(severity, severity)
    run = paragraph.add_run(f'风险等级：{label}')
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string(color)

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('网站安全评估报告')
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor.from_string('1F4E79')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f'{TARGET_DOMAIN} — {TARGET_NAME}')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor.from_string('666666')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f'评估日期：{ASSESSMENT_DATE}')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor.from_string('999999')

doc.add_page_break()

# Target info
if TARGET_INFO:
    doc.add_heading('目标信息', level=1)
    make_table(doc, ['项目', '详情'], TARGET_INFO)

# Risk summary
doc.add_heading('风险总览', level=1)
counts = {}
for sev, *_ in FINDINGS:
    counts[sev] = counts.get(sev, 0) + 1
summary_rows = [[RISK_LABELS[s], str(c), ''] for s, c in counts.items()]
make_table(doc, ['风险等级', '数量', '代表性问题'], summary_rows)

# Detailed findings by severity
for sev in ['critical', 'high', 'medium', 'low']:
    items = [f for f in FINDINGS if f[0] == sev]
    if not items:
        continue
    doc.add_heading(f'{RISK_LABELS[sev]}风险', level=1)
    for i, (_, title, desc, evidence, impact, remediation, priority) in enumerate(items):
        doc.add_heading(f'{i+1}. {title}', level=2)
        p = doc.add_paragraph()
        add_risk_badge(p, sev)
        doc.add_paragraph(desc)
        if evidence:
            doc.add_paragraph(str(evidence), style='Intense Quote')
        if impact:
            p = doc.add_paragraph()
            run = p.add_run('影响：')
            run.bold = True
            doc.add_paragraph(impact)
        if remediation:
            p = doc.add_paragraph()
            run = p.add_run('修复建议：')
            run.bold = True
            doc.add_paragraph(remediation)

# Remediation summary
if FINDINGS:
    doc.add_page_break()
    doc.add_heading('修复建议汇总', level=1)
    rows = [[str(i+1), f[1], RISK_LABELS[f[0]], f[5], f[6]] for i, f in enumerate(FINDINGS)]
    make_table(doc, ['#', '问题', '风险', '修复建议', '优先级'], rows)

doc.save(OUTPUT_PATH)
print(f"报告已保存到: {OUTPUT_PATH}")
