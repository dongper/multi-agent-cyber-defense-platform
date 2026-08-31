#!/usr/bin/env python3
"""
Generate a 京网杯-formatted DOCX from structured content.
Usage: python generate_docx.py <output_path> <content_json>

Content JSON schema:
{
  "title": "论文标题",
  "authors": ["作者1", "作者2"],
  "unit": "单位全称，邮编",
  "abstract": "摘要内容...",
  "keywords": ["关键词1", "关键词2"],
  "sections": [
    {"level": 1, "title": "1 引言", "text": "段落内容..."},
    {"level": 2, "title": "1.1 小节", "text": "段落内容..."},
    {"type": "image", "path": "/path/to/img.png", "caption": "图1 xxx"},
    {"type": "table", "caption": "表1 xxx", "headers": ["col1","col2"], "rows": [["a","b"]]}
  ],
  "references": ["[1] Author. Title[J]. Journal, Year.", ...],
  "bios": ["姓名（出生年月—），性别，职称，学位，单位，专业，邮箱。", ...]
}
"""
import json, sys, os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# Font size mapping (Chinese standard)
SIZES = {
    "小二": Pt(18), "四号": Pt(14), "小四": Pt(12),
    "五号": Pt(10.5), "六号": Pt(9)
}

def set_cn_font(run, name, size, bold=False):
    run.font.size = size
    run.font.bold = bold
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)

def generate_docx(output_path, content):
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2.54); s.bottom_margin = Cm(2.54)
        s.left_margin = Cm(3.17); s.right_margin = Cm(3.17)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cn_font(p.add_run(content["title"]), '黑体', SIZES["小二"], bold=True)

    doc.add_paragraph()

    # Authors
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cn_font(p.add_run("  ".join(content["authors"])), '楷体', SIZES["五号"])

    # Unit
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cn_font(p.add_run(content["unit"]), '宋体', SIZES["五号"])

    doc.add_paragraph()

    # Abstract
    p = doc.add_paragraph()
    set_cn_font(p.add_run("【摘要】"), '宋体', SIZES["五号"], bold=True)
    set_cn_font(p.add_run(content["abstract"]), '宋体', SIZES["五号"])

    # Keywords
    p = doc.add_paragraph()
    set_cn_font(p.add_run("关键词："), '宋体', SIZES["五号"], bold=True)
    set_cn_font(p.add_run("；".join(content["keywords"])), '宋体', SIZES["五号"])

    doc.add_paragraph()

    # Body sections
    for sec in content.get("sections", []):
        if sec.get("type") == "image":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(sec["path"], width=Cm(14))
            if sec.get("caption"):
                cap = doc.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_cn_font(cap.add_run(sec["caption"]), '宋体', SIZES["五号"])

        elif sec.get("type") == "table":
            if sec.get("caption"):
                cap = doc.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_cn_font(cap.add_run(sec["caption"]), '宋体', SIZES["五号"])
            rows = len(sec["rows"]) + 1
            cols = len(sec["headers"])
            tbl = doc.add_table(rows=rows, cols=cols, style='Table Grid')
            tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for i, h in enumerate(sec["headers"]):
                cell = tbl.rows[0].cells[i]
                cell.text = h
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].runs[0].font.size = SIZES["五号"]
            for i, row in enumerate(sec["rows"]):
                for j, val in enumerate(row):
                    cell = tbl.rows[i+1].cells[j]
                    cell.text = str(val)
                    cell.paragraphs[0].runs[0].font.size = SIZES["五号"]
        else:
            # Heading
            level = sec.get("level", 0)
            if level >= 1:
                p = doc.add_paragraph()
                font_size = {1: SIZES["四号"], 2: SIZES["小四"], 3: SIZES["五号"]}.get(level, SIZES["五号"])
                set_cn_font(p.add_run(sec["title"]), '黑体', font_size, bold=True)

            # Body text (supports \n\n for multiple paragraphs)
            if sec.get("text"):
                paragraphs = sec["text"].split("\n\n")
                for para_text in paragraphs:
                    para_text = para_text.strip()
                    if para_text:
                        p = doc.add_paragraph()
                        p.paragraph_format.first_line_indent = Cm(0.74)
                        set_cn_font(p.add_run(para_text), '宋体', SIZES["五号"])

    # References
    if content.get("references"):
        p = doc.add_paragraph()
        set_cn_font(p.add_run("参考文献"), '黑体', SIZES["四号"], bold=True)
        for ref in content["references"]:
            p = doc.add_paragraph()
            set_cn_font(p.add_run(ref), '宋体', SIZES["六号"])

    # Author bios
    if content.get("bios"):
        p = doc.add_paragraph()
        set_cn_font(p.add_run("作者简介"), '宋体', SIZES["五号"], bold=True)
        for bio in content["bios"]:
            p = doc.add_paragraph()
            set_cn_font(p.add_run(bio), '楷体', SIZES["六号"])

    doc.save(output_path)
    return output_path

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python generate_docx.py <output_path> <content_json_file>")
        sys.exit(1)
    with open(sys.argv[2], 'r', encoding='utf-8') as f:
        content = json.load(f)
    print(f"Generated: {generate_docx(sys.argv[1], content)}")
