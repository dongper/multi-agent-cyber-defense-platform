#!/usr/bin/env python3
"""
Fix citations in an existing 京网杯 DOCX file:
1. Delete specified references (by old number)
2. Renumber remaining references
3. Update all in-text citations to match new numbering
4. Set in-text citations to superscript
5. Remove text mentions of deleted papers

Usage:
    python fix_citations.py paper.docx --delete 3,8,9,10

⚠️ ALWAYS backup the file first. This script modifies in-place.
"""

import argparse
import copy
import re
import shutil
from docx import Document
from docx.oxml.ns import qn
from lxml import etree


def build_mapping(total_refs, delete_set):
    """Build old_num -> new_num mapping. Deleted refs map to None."""
    mapping = {}
    new_num = 1
    for old in range(1, total_refs + 1):
        if old in delete_set:
            mapping[old] = None
        else:
            mapping[old] = new_num
            new_num += 1
    return mapping


def delete_ref_paragraphs(doc, delete_set):
    """Remove reference paragraphs for deleted refs (back-to-front)."""
    body = doc.element.body
    ref_start = None
    delete_indices = []
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip() == '参考文献':
            ref_start = i
        if ref_start and i > ref_start:
            m = re.match(r'^\[(\d+)\]', para.text.strip())
            if m and int(m.group(1)) in delete_set:
                delete_indices.append(i)
    for idx in sorted(delete_indices, reverse=True):
        body.remove(doc.paragraphs[idx]._element)
    return ref_start


def renumber_references(doc, mapping):
    """Renumber [N] in reference list paragraphs via lxml."""
    nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    ref_start = None
    for para in doc.paragraphs:
        if para.text.strip() == '参考文献':
            ref_start = para
            continue
        if ref_start is None:
            continue
        text = para.text.strip()
        m = re.match(r'^\[(\d+)\]', text)
        if not m:
            if text and not text.startswith('['):
                break  # past references section
            continue
        old = int(m.group(1))
        new = mapping.get(old)
        if new is None:
            continue
        for t in para._element.findall('.//w:t', nsmap):
            if t.text and re.match(r'^\[\d+\]', t.text):
                t.text = re.sub(r'^\[\d+\]', f'[{new}]', t.text, count=1)


def update_in_text_citations(doc, mapping):
    """Update all [N] in body text using single-pass regex (avoids sequential collision)."""
    nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    ref_start = None
    for para in doc.paragraphs:
        if para.text.strip() == '参考文献':
            break
        for t in para._element.findall('.//w:t', nsmap):
            if t.text and '[' in t.text:
                def repl(m):
                    n = mapping.get(int(m.group(1)))
                    return f'[{n}]' if n else ''
                t.text = re.sub(r'\[(\d+)\]', repl, t.text)


def set_superscript_citations(doc):
    """Set all [N] runs in body text to superscript."""
    ref_start = None
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip() == '参考文献':
            break
        if not re.search(r'\[\d+\]', para.text):
            continue
        # Split runs that contain mixed text and [N]
        new_runs_data = []
        for run in para.runs:
            if not re.search(r'\[\d+\]', run.text):
                new_runs_data.append(('text', run.text, run._element.find(qn('w:rPr'))))
                continue
            parts = re.split(r'(\[\d+\])', run.text)
            rpr = run._element.find(qn('w:rPr'))
            for part in parts:
                if part == '':
                    continue
                is_citation = bool(re.match(r'^\[\d+\]$', part))
                new_runs_data.append(('cite' if is_citation else 'text', part, rpr))

        # Clear all existing runs
        for run in para.runs:
            run.text = ''

        # Rebuild runs
        first_run = para.runs[0]
        first_run.text = new_runs_data[0][1] if new_runs_data else ''
        if new_runs_data and new_runs_data[0][0] == 'cite':
            _set_superscript(first_run._element)

        parent = first_run._element.getparent()
        insert_after = first_run._element

        for kind, text, rpr_template in new_runs_data[1:]:
            new_run = etree.SubElement(parent, qn('w:r'))
            if rpr_template is not None:
                new_rpr = copy.deepcopy(rpr_template)
                new_run.insert(0, new_rpr)
            t = etree.SubElement(new_run, qn('w:t'))
            t.text = text
            t.set(qn('xml:space'), 'preserve')
            if kind == 'cite':
                _set_superscript(new_run)
            insert_after.addnext(new_run)
            insert_after = new_run


def _set_superscript(run_elem):
    """Set a run element to superscript."""
    rpr = run_elem.find(qn('w:rPr'))
    if rpr is None:
        rpr = etree.SubElement(run_elem, qn('w:rPr'))
        run_elem.insert(0, rpr)
    va = rpr.find(qn('w:vertAlign'))
    if va is None:
        va = etree.SubElement(rpr, qn('w:vertAlign'))
    va.set(qn('w:val'), 'superscript')


def add_citation_after_text(para, search_text, citation_num):
    """Insert a superscript [N] after search_text in a paragraph.
    
    Use when a paper mentions a concept (e.g., 'LoRA微调', 'RAG') 
    without a citation, and you need to add one.
    
    Finds the run containing search_text, splits it, and inserts 
    a superscript [citation_num] run between the matched text and 
    the remainder.
    
    Returns True if successful, False if search_text not found.
    """
    for run in para.runs:
        if search_text not in run.text:
            continue
        idx = run.text.index(search_text) + len(search_text)
        before = run.text[:idx]
        after = run.text[idx:]
        run.text = before

        rpr = run._element.find(qn('w:rPr'))

        # Create superscript citation run
        cite_run = etree.SubElement(run._element.getparent(), qn('w:r'))
        if rpr is not None:
            cite_rpr = copy.deepcopy(rpr)
            cite_run.insert(0, cite_rpr)
        else:
            cite_rpr = etree.SubElement(cite_run, qn('w:rPr'))
        _set_superscript(cite_run)
        t = etree.SubElement(cite_run, qn('w:t'))
        t.text = f'[{citation_num}]'
        t.set(qn('xml:space'), 'preserve')
        run._element.addnext(cite_run)

        # Create trailing text run (if any text remains)
        if after:
            after_run = etree.SubElement(run._element.getparent(), qn('w:r'))
            if rpr is not None:
                after_rpr = copy.deepcopy(rpr)
                after_run.insert(0, after_rpr)
            t2 = etree.SubElement(after_run, qn('w:t'))
            t2.text = after
            t2.set(qn('xml:space'), 'preserve')
            cite_run.addnext(after_run)
        return True
    return False


def remove_sentences_containing(doc, keywords, ref_start_idx=None):
    """Remove sentences containing any of the given keywords from body text.
    
    Use after deleting references to clean up text that mentions the 
    deleted papers by author name, paper title, or system name.
    
    Args:
        doc: Document object
        keywords: list of strings to match (case-sensitive)
        ref_start_idx: paragraph index of '参考文献' heading (auto-detected if None)
    
    Returns: number of paragraphs modified
    """
    if ref_start_idx is None:
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip() == '参考文献':
                ref_start_idx = i
                break
    
    modified = 0
    for i, para in enumerate(doc.paragraphs):
        if ref_start_idx and i >= ref_start_idx:
            break
        text = para.text
        if not any(kw in text for kw in keywords):
            continue
        
        # Split into sentences (Chinese sentence endings)
        sentences = re.split(r'(?<=[。！？])', text)
        new_sentences = [s for s in sentences if not any(kw in s for kw in keywords)]
        new_text = ''.join(new_sentences)
        
        if new_text != text:
            # Clear all runs, set first run to new text
            for run in para.runs:
                run.text = ''
            if para.runs:
                para.runs[0].text = new_text
            modified += 1
    
    return modified


def main():
    parser = argparse.ArgumentParser(description='Fix citations in 京网杯 DOCX')
    parser.add_argument('file', help='DOCX file to edit')
    parser.add_argument('--delete', required=True, help='Comma-separated old ref numbers to delete')
    parser.add_argument('--total', type=int, default=15, help='Total number of references (default: 15)')
    parser.add_argument('--no-superscript', action='store_true', help='Skip superscript formatting')
    parser.add_argument('--no-backup', action='store_true', help='Skip backup (dangerous)')
    args = parser.parse_args()

    delete_set = set(int(x) for x in args.delete.split(','))
    mapping = build_mapping(args.total, delete_set)

    print(f"Mapping: { {k: v for k, v in mapping.items() if v is not None} }")
    print(f"Deleting: {delete_set}")

    # Backup
    if not args.no_backup:
        bak = args.file + '.bak'
        shutil.copy2(args.file, bak)
        print(f"Backup: {bak}")

    doc = Document(args.file)

    # Phase 1: Delete ref paragraphs
    ref_start = delete_ref_paragraphs(doc, delete_set)
    print(f"Deleted {len(delete_set)} reference paragraphs")

    # Phase 2: Renumber references
    renumber_references(doc, mapping)

    # Phase 3: Update in-text citations
    update_in_text_citations(doc, mapping)

    # Phase 4: Superscript
    if not args.no_superscript:
        set_superscript_citations(doc)
        print("Set in-text citations to superscript")

    doc.save(args.file)
    print(f"Saved: {args.file}")

    # Verify
    doc2 = Document(args.file)
    all_citations = set()
    for para in doc2.paragraphs:
        if para.text.strip() == '参考文献':
            break
        for r in re.findall(r'\[(\d+)\]', para.text):
            all_citations.add(int(r))
    
    expected = set(mapping[k] for k, v in mapping.items() if v is not None)
    missing = expected - all_citations
    extra = all_citations - expected
    if missing:
        print(f"⚠️ Orphan refs (in list, not cited): {missing}")
    if extra:
        print(f"⚠️ Dangling citations (cited, no ref): {extra}")
    if not missing and not extra:
        print(f"✓ All {len(expected)} references cited correctly")


if __name__ == '__main__':
    main()
