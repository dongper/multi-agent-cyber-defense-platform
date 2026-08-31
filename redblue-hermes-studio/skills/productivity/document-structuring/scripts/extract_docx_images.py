"""Extract images and paragraph map from a DOCX file.

Usage:
    python extract_docx_images.py input.docx [output_dir]

Outputs:
    - Extracted images to output_dir (default: ./extracted_images/)
    - Paragraph-to-image mapping printed to stdout
    - Paragraph text inventory printed to stdout
"""
import sys, os
from docx import Document
from docx.oxml.ns import qn

def main():
    if len(sys.argv) < 2:
        print("Usage: python extract_docx_images.py input.docx [output_dir]")
        sys.exit(1)

    docx_path = sys.argv[1]
    out_dir = sys.argv[2] if len(sys.argv) > 2 else './extracted_images'
    os.makedirs(out_dir, exist_ok=True)

    doc = Document(docx_path)
    rels = doc.part.rels

    # 1. Paragraph text inventory
    print("=== PARAGRAPHS ===")
    for i, para in enumerate(doc.paragraphs):
        text = para.text[:120] if para.text else ''
        if text.strip():
            print(f'[{i}] style={para.style.name} | {text}')

    # 2. Image inventory
    print("\n=== IMAGES IN DOC ===")
    image_rels = {}
    for rel_id, rel in rels.items():
        if 'image' in rel.reltype:
            fname = rel.target_ref.split('/')[-1]
            image_rels[rel_id] = fname
            fpath = os.path.join(out_dir, fname)
            with open(fpath, 'wb') as f:
                f.write(rel.target_part.blob)
            print(f'  {rel_id}: {fname} ({len(rel.target_part.blob)} bytes) -> saved')

    # 3. Image-to-paragraph mapping
    print("\n=== IMAGE PLACEMENT ===")
    for i, para in enumerate(doc.paragraphs):
        drawings = para._element.findall('.//' + qn('w:drawing'))
        for d in drawings:
            blips = d.findall('.//' + qn('a:blip'))
            for blip in blips:
                embed = blip.get(qn('r:embed'))
                if embed in image_rels:
                    print(f'  Para [{i}] text="{para.text[:60]}" -> {image_rels[embed]}')

    # 4. Tables
    print(f"\n=== TABLES ({len(doc.tables)}) ===")
    for ti, table in enumerate(doc.tables):
        print(f'Table {ti}: {len(table.rows)}r x {len(table.columns)}c')
        for ri, row in enumerate(table.rows):
            cells = [c.text[:30].replace('\n',' ') for c in row.cells]
            print(f'  Row {ri}: {cells}')
            if ri > 10:
                print('  ... (truncated)')
                break

if __name__ == '__main__':
    main()
